import re
import numpy as np
from transformers import pipeline, AutoTokenizer, AutoModelForCausalLM
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import torch
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from datetime import datetime
import os

DEDUP_SIMILARITY_THRESHOLD = 0.85

MIN_REQUIREMENT_LENGTH = 25
NOISE_PHRASES = [
    "hi team", "after our meeting", "regards", "thanks", "thank you",
    "please find", "as discussed", "following up", "hope this", "let me know",
    "subject:", "date:", "from:", "to:", "cc:", "best regards", "kind regards",
    "dear ", "sincerely", "attached please", "per our", "as per",
]

REQUIREMENT_KEYWORDS = [
    "must", "should", "shall", "require", "needs to", "need to",
    "will", "expected to", "has to", "have to",
]

STAKEHOLDER_MAPPING = {
    "user":           "Users",
    "users":          "Users",
    "admin":          "Admin",
    "administrator":  "Admin",
    "administrators": "Admin",
    "stakeholder":    "Business Stakeholder",
    "stakeholder 2":  "Business Stakeholder",
    "team":           "Internal Team",
    "manager":        "Product Manager",
    "product manager":"Product Manager",
    "ux designer":    "UX / Design Team",
    "designer":       "UX / Design Team",
    "general":        "General",
}


class BRDGenerator:
    """
    ML-powered Business Requirements Document generator.

    Fixes applied (v3):
      1. Truncated titles  → always derived from full description, never req.title
      2. Duplicate reqs    → semantic deduplication via sentence embeddings
      3. Garbage entries   → strict validation gate before any req is processed
      4. Stakeholder dupes → canonical name normalisation
      5. NFR misclassify   → zero-shot classification for req type
      6. Hallucinated URLs → constrained executive-summary prompt
    """

    def __init__(self):
        self.summarizer = pipeline(
            "summarization",
            model="facebook/bart-large-cnn",
        )

        self.sentence_model = SentenceTransformer("all-MiniLM-L6-v2")

        self.generator = pipeline(
            "text2text-generation",
            model="google/flan-t5-base",
        )

        # Fix #5 — zero-shot classifier for requirement type correction.
        self.classifier = pipeline(
            "zero-shot-classification",
            model="facebook/bart-large-mnli",
        )

    def is_valid_requirement(self, text: str) -> bool:
        """
        Return True only if text looks like a real requirement sentence.
        Rejects metadata, email body noise, greetings, and stubs.
        """
        if not text or len(text.strip()) < MIN_REQUIREMENT_LENGTH:
            return False

        lower = text.lower()

        # Reject if it contains known noise phrases
        if any(phrase in lower for phrase in NOISE_PHRASES):
            return False

        # Reject if it contains no action verb typical of requirements
        if not any(kw in lower for kw in REQUIREMENT_KEYWORDS):
            return False

        return True

    # =======================================================================
    # FIX #4 — STAKEHOLDER NORMALISATION
    # =======================================================================

    def normalize_stakeholder(self, raw: str) -> str:
        """
        Map raw stakeholder labels (user, Users, admin, Stakeholder 2, …)
        to canonical display names.
        """
        if not raw:
            return "General"
        normalised = STAKEHOLDER_MAPPING.get(raw.strip().lower(), raw.strip().title())
        return normalised

    # =======================================================================
    # FIX #5 — REQUIREMENT TYPE CORRECTION VIA ZERO-SHOT CLASSIFICATION
    # =======================================================================

    # Sentences matching these keywords are almost certainly non-functional
    # regardless of how they were tagged in the DB. We use this as a fast-path
    # before calling the expensive zero-shot model.
    _NFR_KEYWORDS = [
        "load pages in under", "response time", "concurrent users",
        "uptime", "availability", "latency", "throughput", "scalab",
        "encrypt", "multi-factor", "mfa", "ssl", "tls", "gdpr",
        "recovery time", "backup", "disaster", "accessibility", "wcag",
    ]

    def correct_requirement_type(self, req) -> str:
        """
        Return the best-guess requirement type for req:
        'functional' | 'non_functional' | 'business'

        Uses a fast keyword check first; falls back to zero-shot classification
        only when the keyword check is inconclusive.
        """
        desc_lower = req.description.lower()

        # Fast path — clear NFR signals
        if any(kw in desc_lower for kw in self._NFR_KEYWORDS):
            return "non_functional"

        # If already non-functional or business and passes keyword check, trust it
        if req.requirement_type in ("non_functional", "business"):
            return req.requirement_type

        # Zero-shot classification for ambiguous cases
        try:
            labels = ["functional requirement", "non-functional requirement", "business requirement"]
            result = self.classifier(req.description[:512], candidate_labels=labels)
            top_label = result["labels"][0]
            if "non-functional" in top_label:
                return "non_functional"
            elif "business" in top_label:
                return "business"
            else:
                return "functional"
        except Exception:
            return req.requirement_type  # fall back to stored value

    # =======================================================================
    # FIX #1 — TITLE DERIVATION (never use req.title)
    # =======================================================================

    def _clean_requirement_text(self, text: str) -> str:
        """
        Strip email headers, speaker labels, domain fragments, truncation
        artefacts, and leading list numbers from raw requirement text.
        """
        if not text:
            return ""

        # Email header lines
        text = re.sub(r"^(From|To|CC|Subject|Date)\s*:.*$", "", text, flags=re.MULTILINE)

        # Speaker / role labels at line start, e.g. "Product Manager: "
        text = re.sub(r"^[A-Z][a-zA-Z ]+:\s*", "", text, flags=re.MULTILINE)

        # Bare domain fragments on their own line
        text = re.sub(r"^\s*(com|org|net|co)\s*$", "", text, flags=re.MULTILINE)

        # Trailing ellipsis truncation artefacts
        text = re.sub(r"\.\.\.\s*$", "", text.strip())

        # Leading list numbering artefacts "1. "
        text = re.sub(r"^\d+\.\s*", "", text.strip())

        # Collapse excess blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()

    def _title_from_description(self, description: str) -> str:
        """
        Derive a clean, complete (non-truncated) display title from the full
        description text.  Never touches req.title.

        1. Clean with regex.
        2. Take the first sentence (split on '.', '?', '!').
        3. If still >120 chars, compress with flan-t5.
        """
        cleaned = self._clean_requirement_text(description)
        if not cleaned:
            return ""

        # First sentence
        match = re.match(r"^([^.?!]+[.?!]?)", cleaned)
        title = match.group(1).strip() if match else cleaned

        if len(title) > 120:
            title = self._generate_title_with_model(cleaned)

        return title

    def _generate_title_with_model(self, description: str) -> str:
        """
        Ask flan-t5 to produce a concise (≤12 word) professional requirement
        title from a longer description.
        """
        prompt = (
            "Generate a short professional requirement title (maximum 12 words) "
            "for the following requirement. "
            "Do not include email headers, speaker names, or metadata.\n"
            f"Requirement: {description[:300]}\n"
            "Title:"
        )
        try:
            result = self.generator(prompt, max_length=30, do_sample=False)
            return result[0]["generated_text"].strip()
        except Exception:
            # Graceful fallback: first 100 chars of cleaned text
            return self._clean_requirement_text(description)[:100]


    def _deduplicate_requirements(self, requirements: list) -> list:
        """
        Remove semantically duplicate requirements using cosine similarity on
        sentence embeddings of the cleaned description field.

        Returns a plain Python list of unique requirement objects.
        The first occurrence of each group is kept (stable order).
        """
        if not requirements:
            return []

        texts = [
            self._clean_requirement_text(r.description) or r.description
            for r in requirements
        ]

        embeddings = self.sentence_model.encode(texts)
        sim_matrix = cosine_similarity(embeddings)

        seen: set = set()
        unique: list = []

        for i, req in enumerate(requirements):
            if i in seen:
                continue
            unique.append(req)
            for j in range(i + 1, len(requirements)):
                if j not in seen and sim_matrix[i, j] >= DEDUP_SIMILARITY_THRESHOLD:
                    seen.add(j)

        return unique


    def generate_brd(self, project, include_conflicts=True,
                     include_traceability=True, include_sentiment=True):
        """
        Generate a complete BRD.

        Pipeline:
          raw DB requirements
            → validation gate          (fix #3)
            → type correction          (fix #5)
            → stakeholder normalisation (fix #4)
            → deduplication            (fix #2)
            → section generation       (fixes #1, #6)
        """
        raw_requirements = list(project.requirements.all())
        data_sources = project.data_sources.filter(is_relevant=True)
        valid_requirements = [
            r for r in raw_requirements
            if self.is_valid_requirement(r.description)
        ]
        for req in valid_requirements:
            req._corrected_type = self.correct_requirement_type(req)
        for req in valid_requirements:
            req._normalised_stakeholder = self.normalize_stakeholder(req.stakeholder)
        unique_requirements = self._deduplicate_requirements(valid_requirements)

        brd_content = {
            "title": f"Business Requirements Document - {project.name}",
            "executive_summary": self._generate_executive_summary(project, unique_requirements),
            "business_objectives": self._generate_business_objectives(unique_requirements),
            "stakeholder_analysis": self._generate_stakeholder_analysis(unique_requirements, data_sources),
            "functional_requirements": self._generate_functional_requirements(unique_requirements),
            "non_functional_requirements": self._generate_non_functional_requirements(unique_requirements),
            "assumptions": self._generate_assumptions(unique_requirements, data_sources),
            "success_metrics": self._generate_success_metrics(unique_requirements),
            "timeline": self._generate_timeline(unique_requirements, data_sources),
        }

        if include_conflicts:
            brd_content["conflict_analysis"] = self._generate_conflict_analysis(project)

        if include_traceability:
            brd_content["traceability_matrix"] = self._generate_traceability_matrix(unique_requirements)

        if include_sentiment:
            brd_content["sentiment_analysis"] = self._analyze_sentiment(data_sources)

        return brd_content

    def _filter_by_type(self, requirements: list, req_type: str) -> list:
        return [r for r in requirements if getattr(r, "_corrected_type", r.requirement_type) == req_type]

    def _filter_by_priority(self, requirements: list, priority: str) -> list:
        return [r for r in requirements if r.priority == priority]

    def _filter_by_stakeholder(self, requirements: list, stakeholder_display: str) -> list:
        return [
            r for r in requirements
            if getattr(r, "_normalised_stakeholder", self.normalize_stakeholder(r.stakeholder)) == stakeholder_display
        ]

    def _generate_executive_summary(self, project, requirements: list) -> str:
        """
        Fix #6: constrained prompt — model is explicitly told not to invent
        URLs, links, or any information not present in the project data.
        """
        total_reqs = len(requirements)
        functional_reqs  = len(self._filter_by_type(requirements, "functional"))
        nfr_count        = len(self._filter_by_type(requirements, "non_functional"))
        business_count   = len(self._filter_by_type(requirements, "business"))

        summary_prompt = (
            f"Write a professional executive summary for a Business Requirements Document.\n\n"
            f"Rules:\n"
            f"- Do NOT invent URLs, links, or website addresses.\n"
            f"- Do NOT add any external references.\n"
            f"- Only use the project information provided below.\n"
            f"- Write 2–3 sentences maximum.\n\n"
            f"Project Name: {project.name}\n"
            f"Description: {project.description}\n"
            f"Total Requirements: {total_reqs}\n"
            f"Functional: {functional_reqs} | Non-Functional: {nfr_count} | Business: {business_count}\n\n"
            f"Executive Summary:"
        )

        try:
            summary = self.summarizer(
                summary_prompt,
                max_length=150,
                min_length=40,
                do_sample=False,
            )[0]["summary_text"]
        except Exception:
            summary = (
                f"This document defines the business requirements for {project.name}. "
                f"It covers {total_reqs} requirements gathered from stakeholder communications "
                f"and project documentation."
            )

        return (
            "EXECUTIVE SUMMARY\n\n"
            f"{summary}\n\n"
            "This Business Requirements Document (BRD) provides a comprehensive overview of the\n"
            "requirements gathered from multiple sources including emails, meeting transcripts,\n"
            "and stakeholder communications.\n\n"
            f"Total Requirements Identified: {total_reqs}\n"
            f"- Functional Requirements: {functional_reqs}\n"
            f"- Non-Functional Requirements: {nfr_count}\n"
            f"- Business Requirements: {business_count}\n\n"
            f"Document Version: 1.0\n"
            f"Date: {datetime.now().strftime('%B %d, %Y')}"
        ).strip()

    def _generate_business_objectives(self, requirements: list) -> str:
        business_reqs = self._filter_by_type(requirements, "business")

        section = "BUSINESS OBJECTIVES\n\n"
        section += "The following business objectives have been identified:\n\n"

        if not business_reqs:
            section += "Business objectives will be derived from the functional requirements.\n"
            return section

        for idx, req in enumerate(business_reqs, 1):
            title       = self._title_from_description(req.description)
            description = self._clean_requirement_text(req.description)
            stakeholder = getattr(req, "_normalised_stakeholder", self.normalize_stakeholder(req.stakeholder))

            section += f"{idx}. {title}\n"
            section += f"   {description}\n"
            section += f"   Priority: {req.priority.upper()}\n"
            section += f"   Stakeholder: {stakeholder}\n"
            section += f"   Source: {req.data_source.source_type}\n\n"

        return section

    def _generate_stakeholder_analysis(self, requirements: list, data_sources) -> str:
        seen_names: set = set()
        stakeholders: list = []
        for req in requirements:
            name = getattr(req, "_normalised_stakeholder", self.normalize_stakeholder(req.stakeholder))
            if name not in seen_names:
                seen_names.add(name)
                stakeholders.append(name)
        stakeholders.sort()

        analysis = "STAKEHOLDER ANALYSIS\n\n"
        analysis += f"Total Stakeholders Identified: {len(stakeholders)}\n\n"

        for stakeholder in stakeholders:
            stakeholder_reqs = self._filter_by_stakeholder(requirements, stakeholder)
            concerns = [
                self._title_from_description(r.description)
                for r in stakeholder_reqs[:3]
            ]
            concerns = [c for c in concerns if c]

            analysis += f"Stakeholder: {stakeholder}\n"
            analysis += f"- Related Requirements: {len(stakeholder_reqs)}\n"
            analysis += f"- Primary Concerns: {', '.join(concerns) if concerns else 'N/A'}\n\n"

        return analysis

    def _generate_functional_requirements(self, requirements: list) -> str:
        """Fix #1: title always from description. Fix #5: type-corrected list."""
        functional_reqs = self._filter_by_type(requirements, "functional")

        section  = "FUNCTIONAL REQUIREMENTS\n\n"
        section += "The system shall provide the following functional capabilities:\n\n"

        for priority in ["high", "medium", "low"]:
            priority_reqs = self._filter_by_priority(functional_reqs, priority)
            if not priority_reqs:
                continue

            section += f"\n{priority.upper()} PRIORITY:\n\n"
            for idx, req in enumerate(priority_reqs, 1):
                title       = self._title_from_description(req.description)   # Fix #1
                description = self._clean_requirement_text(req.description)
                stakeholder = getattr(req, "_normalised_stakeholder", self.normalize_stakeholder(req.stakeholder))

                section += f"FR-{priority[0].upper()}{idx}: {title}\n"
                section += f"Description: {description}\n"
                section += f"Stakeholder: {stakeholder}\n"
                section += f"Confidence: {req.confidence_score:.2f}\n"
                section += f"Source: {req.data_source.source_identifier}\n\n"

        return section

    def _generate_non_functional_requirements(self, requirements: list) -> str:
        """Fix #5: uses type-corrected list, so misclassified NFRs now appear here."""
        nfr = self._filter_by_type(requirements, "non_functional")

        section = "NON-FUNCTIONAL REQUIREMENTS\n\n"

        categories = {
            "Performance":  ["performance", "speed", "response", "load", "concurrent", "latency", "throughput"],
            "Security":     ["security", "authentication", "authorization", "encryption", "multi-factor", "mfa", "ssl", "tls", "gdpr"],
            "Usability":    ["usability", "user interface", "intuitive", "experience", "responsive", "accessibility", "wcag"],
            "Reliability":  ["reliability", "availability", "uptime", "recovery", "backup", "disaster"],
            "Scalability":  ["scalability", "scale", "growth", "elastic"],
        }

        found_any = False
        for category, keywords in categories.items():
            cat_reqs = [
                r for r in nfr
                if any(kw in r.description.lower() for kw in keywords)
            ]
            if not cat_reqs:
                continue

            found_any = True
            section += f"\n{category}:\n"
            for idx, req in enumerate(cat_reqs, 1):
                title       = self._title_from_description(req.description)
                description = self._clean_requirement_text(req.description)
                section += f"NFR-{category[:3].upper()}{idx}: {title}\n"
                section += f"   {description}\n\n"

        if not found_any:
            section += "No non-functional requirements have been identified at this time.\n"

        return section

    def _generate_assumptions(self, requirements: list, data_sources) -> str:
        assumptions  = "ASSUMPTIONS AND CONSTRAINTS\n\n"
        assumptions += "The following assumptions have been made:\n\n"
        assumptions += "1. All stakeholders will be available for requirement validation\n"
        assumptions += "2. Required technical infrastructure is available\n"
        assumptions += "3. Budget and resources are allocated as per project plan\n"
        assumptions += f"4. Data gathered from {data_sources.count()} sources is accurate and complete\n"
        assumptions += "5. Requirements may evolve during the project lifecycle\n"
        return assumptions

    def _generate_success_metrics(self, requirements: list) -> str:
        total         = len(requirements)
        high_priority = len(self._filter_by_priority(requirements, "high"))

        metrics  = "SUCCESS METRICS\n\n"
        metrics += "The following metrics will be used to measure project success:\n\n"
        metrics += "1. Requirements Coverage\n"
        metrics += f"   - Total unique requirements: {total}\n"
        metrics += f"   - High priority requirements: {high_priority}\n\n"
        metrics += "2. Stakeholder Satisfaction\n"
        metrics += "   - Measured through feedback surveys and acceptance testing\n\n"
        metrics += "3. Quality Metrics\n"
        metrics += "   - Defect density\n"
        metrics += "   - Test coverage\n"
        metrics += "   - Performance benchmarks\n\n"
        return metrics

    def _generate_timeline(self, requirements: list, data_sources) -> str:
        weeks = max(4, len(requirements) // 5)

        timeline  = "PROJECT TIMELINE\n\n"
        timeline += "Estimated timeline based on requirements complexity:\n\n"
        timeline += "Phase 1: Requirements Analysis (2 weeks)\n"
        timeline += "Phase 2: Design (3 weeks)\n"
        timeline += f"Phase 3: Development ({weeks} weeks)\n"
        timeline += "Phase 4: Testing (2 weeks)\n"
        timeline += "Phase 5: Deployment (1 week)\n\n"
        timeline += f"Total Estimated Duration: {weeks + 8} weeks\n"
        return timeline

    def _generate_conflict_analysis(self, project) -> str:
        conflicts = project.conflicts.all()

        analysis = "CONFLICT ANALYSIS\n\n"
        if conflicts.exists():
            analysis += f"Total Conflicts Detected: {conflicts.count()}\n\n"
            for conflict in conflicts:
                analysis += f"Conflict Type: {conflict.conflict_type}\n"
                analysis += f"Description: {conflict.description}\n"
                analysis += f"Severity: {conflict.severity}\n"
                analysis += f"Status: {'Resolved' if conflict.resolved else 'Pending'}\n\n"
        else:
            analysis += "No conflicts detected between requirements.\n"
        return analysis

    def _generate_traceability_matrix(self, requirements: list) -> dict:
        return {
            f"{getattr(r, '_corrected_type', r.requirement_type).upper()}-{r.id}": {
                "title":       self._title_from_description(r.description),
                "source":      r.data_source.source_type,
                "source_id":   r.data_source.source_identifier,
                "stakeholder": getattr(r, "_normalised_stakeholder", self.normalize_stakeholder(r.stakeholder)),
                "priority":    r.priority,
            }
            for r in requirements
        }

    def _analyze_sentiment(self, data_sources) -> dict:
        sentiment_analyzer = pipeline("sentiment-analysis")
        sentiments = {"positive": 0, "negative": 0, "neutral": 0}

        for source in data_sources[:50]:
            try:
                result = sentiment_analyzer(source.raw_content[:512])[0]
                label  = result["label"].lower()
                if "pos" in label:
                    sentiments["positive"] += 1
                elif "neg" in label:
                    sentiments["negative"] += 1
                else:
                    sentiments["neutral"] += 1
            except Exception:
                sentiments["neutral"] += 1

        return sentiments

    # =======================================================================
    # EDITING & EXPORT
    # =======================================================================

    def apply_edit(self, current_content: str, instruction: str) -> str:
        prompt = (
            f"Edit the following text according to this instruction: {instruction}\n\n"
            f"Text: {current_content}\n\n"
            f"Edited text:"
        )
        try:
            result = self.generator(prompt, max_length=512, num_return_sequences=1)
            return result[0]["generated_text"]
        except Exception:
            return current_content

    def generate_document_file(self, brd) -> str:
        """Render the BRD model instance to a .docx file and return the path."""
        doc = Document()

        heading = doc.add_heading(brd.title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info = doc.add_paragraph()
        info.add_run(f"Version: {brd.version}\n").bold = True
        info.add_run(f"Status: {brd.status.upper()}\n").bold = True
        info.add_run(f"Date: {brd.created_at.strftime('%B %d, %Y')}\n").bold = True

        doc.add_page_break()

        sections = [
            ("Executive Summary",        brd.executive_summary),
            ("Business Objectives",      brd.business_objectives),
            ("Stakeholder Analysis",     brd.stakeholder_analysis),
            ("Functional Requirements",  brd.functional_requirements),
            ("Non-Functional Requirements", brd.non_functional_requirements),
            ("Assumptions and Constraints", brd.assumptions),
            ("Success Metrics",          brd.success_metrics),
            ("Timeline",                 brd.timeline),
        ]

        if brd.conflict_analysis:
            sections.append(("Conflict Analysis", brd.conflict_analysis))

        for section_title, content in sections:
            doc.add_heading(section_title, 1)
            doc.add_paragraph(content)
            doc.add_paragraph()

        filename = f"BRD_{brd.project.name.replace(' ', '_')}_v{brd.version}.docx"
        filepath = os.path.join(
            "/home/claude/brd_generator_project/backend/media", filename
        )
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        doc.save(filepath)
        return filepath